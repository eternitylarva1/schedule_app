"""Note HTTP endpoints."""
import io
import json
from aiohttp import web
from typing import Any
from .. import db
from ..models import Note, NoteGroup
from ._helpers import (
    json_response, error_response, _sanitize_ai_provider,
    _update_note_stats, _handle_note_operation,
)


# ============= Note Handlers =============

"""GET /api/notes/{note_id}/conversations - get note conversations."""
async def get_note_conversations(request: web.Request) -> web.Response:
    """GET /api/notes/{note_id}/conversations - get conversation history for a note."""
    note_id = int(request.match_info["note_id"])
    try:
        conversations = await db.get_note_conversations(note_id)
        return json_response([c.to_dict() for c in conversations])
    except Exception as e:
        return error_response(f"获取对话历史失败: {str(e)}")




"""POST /api/notes/{note_id}/chat - chat with note."""
async def chat_note(request: web.Request) -> web.Response:
    """POST /api/notes/{note_id}/chat - chat with AI about a note."""
    note_id = int(request.match_info["note_id"])
    
    try:
        body_bytes = await request.read()
        try:
            body_str = body_bytes.decode('utf-8')
        except UnicodeDecodeError:
            body_str = body_bytes.decode('gbk', errors='replace')
        data = json.loads(body_str)
    except Exception as e:
        return error_response("无效的JSON数据")
    
    try:
        # Get the note content
        note = await db.get_note(note_id)
        if not note:
            return error_response("笔记不存在", code=404)
        
        user_message = data.get("message", "").strip()
        if not user_message:
            return error_response("消息内容不能为空")
        
        selected_text = data.get("selected_text", "")
        
        # Get conversation history for context
        conversations = await db.get_note_conversations(note_id)
        history_str = ""
        if conversations:
            history_lines = []
            for c in conversations:
                role_label = "用户" if c.role == "user" else "AI"
                history_lines.append(f"{role_label}：{c.content}")
            history_str = "\n".join(history_lines)
        
        # Call LLM service
        from ..llm_service import llm_service
        ai_response = await llm_service.chat_about_note(
            note_content=note.content,
            user_message=user_message,
            selected_text=selected_text,
            conversation_history=history_str
        )
        
        if not ai_response:
            return error_response(llm_service.last_error_message or "AI 响应失败，请重试")
        
        # Save user message to history
        user_conv = db.NoteConversation(
            note_id=note_id,
            role="user",
            content=user_message,
            selected_text=selected_text or ""
        )
        await db.create_note_conversation(user_conv)
        
        # Save AI response to history
        ai_conv = db.NoteConversation(
            note_id=note_id,
            role="assistant",
            content=ai_response,
            selected_text=""
        )
        saved_ai = await db.create_note_conversation(ai_conv)
        
        return json_response(saved_ai.to_dict())
    except Exception as e:
        return error_response(f"AI 对话失败: {str(e)}")




"""DELETE /api/notes/{note_id}/conversations - delete note conversations."""
async def delete_note_conversations(request: web.Request) -> web.Response:
    """DELETE /api/notes/{note_id}/conversations - clear conversation history for a note."""
    note_id = int(request.match_info["note_id"])
    try:
        await db.delete_note_conversations(note_id)
        return json_response({"deleted": True})
    except Exception as e:
        return error_response(f"清空对话历史失败: {str(e)}")


# ============ Expenses Endpoints ============



"""GET /api/notes - list notes."""
async def get_notes(request: web.Request) -> web.Response:
    """GET /api/notes?include_archived=true|false - list all notes."""
    try:
        include_archived_str = request.query.get("include_archived", "false").lower()
        include_archived = include_archived_str in ("true", "1")
        notes = await db.get_notes(include_archived=include_archived)
        return json_response([n.to_dict() for n in notes])
    except Exception as e:
        return error_response(f"获取笔记失败: {str(e)}")




"""POST /api/notes - create note."""
async def create_note(request: web.Request) -> web.Response:
    """POST /api/notes - create a new note."""
    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")
    
    try:
        note = Note(
            title=data.get("title", ""),
            content=data.get("content", ""),
            group_id=data.get("group_id"),
            is_pinned=bool(data.get("is_pinned", False)),
            color=data.get("color", ""),
            is_archived=bool(data.get("is_archived", False)),
        )
        note = await db.create_note(note)
        return json_response(note.to_dict())
    except Exception as e:
        return error_response(f"创建笔记失败: {str(e)}")




"""PUT /api/notes/{id} - update note."""
async def update_note(request: web.Request) -> web.Response:
    """PUT /api/notes/{id} - update a note."""
    note_id = int(request.match_info["id"])
    
    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")
    
    try:
        existing = await db.get_note(note_id)
        if not existing:
            return error_response("笔记不存在", code=404)
        
        # Handle sort_order - use new value if provided, otherwise keep existing
        sort_order = data.get("sort_order")
        if sort_order is None:
            sort_order = existing.sort_order
        
        # Handle new fields - use new value if provided, otherwise keep existing
        is_pinned = data.get("is_pinned")
        if is_pinned is None:
            is_pinned = existing.is_pinned
        
        color = data.get("color")
        if color is None:
            color = existing.color
        
        is_archived = data.get("is_archived")
        if is_archived is None:
            is_archived = existing.is_archived
        
        note = Note(
            title=data.get("title", existing.title),
            content=data.get("content", existing.content),
            group_id=data.get("group_id", existing.group_id),
            sort_order=sort_order,
            is_pinned=is_pinned,
            color=color,
            is_archived=is_archived,
        )
        result = await db.update_note(note_id, note)
        if not result:
            return error_response("笔记不存在", code=404)
        return json_response(result.to_dict())
    except Exception as e:
        return error_response(f"更新笔记失败: {str(e)}")




"""DELETE /api/notes/{id} - delete note."""
async def delete_note(request: web.Request) -> web.Response:
    """DELETE /api/notes/{id} - delete a note."""
    note_id = int(request.match_info["id"])
    
    try:
        success = await db.delete_note(note_id)
        if not success:
            return error_response("笔记不存在", code=404)
        return json_response({"success": True})
    except Exception as e:
        return error_response(f"删除笔记失败: {str(e)}")


# ============================================
# Note Groups API
# ============================================



"""GET /api/note-groups - list note groups."""
async def get_note_groups(request: web.Request) -> web.Response:
    """GET /api/note-groups - list all note groups."""
    try:
        groups = await db.get_note_groups()
        return json_response([g.to_dict() for g in groups])
    except Exception as e:
        return error_response(f"获取笔记分组失败: {str(e)}")




"""POST /api/note-groups - create note group."""
async def create_note_group(request: web.Request) -> web.Response:
    """POST /api/note-groups - create a new note group."""
    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")
    
    try:
        note_group = NoteGroup(
            name=data.get("name", ""),
            sort_order=data.get("sort_order", 0),
        )
        note_group = await db.create_note_group(note_group)
        return json_response(note_group.to_dict())
    except Exception as e:
        return error_response(f"创建笔记分组失败: {str(e)}")




"""PUT /api/note-groups/{id} - update note group."""
async def update_note_group(request: web.Request) -> web.Response:
    """PUT /api/note-groups/{id} - update a note group."""
    group_id = int(request.match_info["id"])
    
    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")
    
    try:
        note_group = NoteGroup(
            name=data.get("name", ""),
            sort_order=data.get("sort_order", 0),
        )
        result = await db.update_note_group(group_id, note_group)
        if not result:
            return error_response("笔记分组不存在", code=404)
        return json_response(result.to_dict())
    except Exception as e:
        return error_response(f"更新笔记分组失败: {str(e)}")




"""DELETE /api/note-groups/{id} - delete note group."""
async def delete_note_group(request: web.Request) -> web.Response:
    """DELETE /api/note-groups/{id} - delete a note group."""
    group_id = int(request.match_info["id"])
    
    try:
        success = await db.delete_note_group(group_id)
        if not success:
            return error_response("笔记分组不存在", code=404)
        return json_response({"success": True})
    except Exception as e:
        return error_response(f"删除笔记分组失败: {str(e)}")


"""PUT /api/notes/reorder - batch reorder notes."""
async def reorder_notes(request: web.Request) -> web.Response:
    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")
    note_ids = data.get("note_ids")
    if not note_ids or not isinstance(note_ids, list):
        return error_response("缺少 note_ids 数组")
    try:
        await db.reorder_notes([int(n) for n in note_ids])
        return json_response({"success": True, "count": len(note_ids)})
    except Exception as e:
        return error_response(f"排序保存失败: {str(e)}")


# ============= Note Export =============

NOTO_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"

def _html_to_text(html: str) -> str:
    """Strip HTML tags and unescape entities for plain-text export."""
    import re, html as html_mod
    text = re.sub(r'<br\s*/?>|</p>|</div>', '\n', html or '', flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    return html_mod.unescape(text)


def _export_pdf_bytes(title: str, content: str) -> bytes:
    """Generate a PDF file in memory using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_font("Noto", fname=NOTO_FONT_PATH, collection_font_number=2)
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.add_page()
    # Title
    pdf.set_font("Noto", size=18)
    pdf.multi_cell(0, 10, title or "未命名笔记")
    pdf.ln(4)
    # Separator line
    w = pdf.w - 2 * pdf.l_margin
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + w, pdf.get_y())
    pdf.ln(6)
    # Content — split paragraphs
    pdf.set_font("Noto", size=11)
    for line in content.splitlines():
        if line.strip():
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 7, line)
        else:
            pdf.ln(4)

    # Output to bytes
    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.getvalue()


def _export_docx_bytes(title: str, content: str) -> bytes:
    """Generate a .docx file in memory using python-docx."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Adjust default style
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # Title
    heading = doc.add_heading(title or "未命名笔记", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Content — split paragraphs
    for line in content.splitlines():
        paragraph = doc.add_paragraph(line if line.strip() else " ")
        paragraph.style.font.size = Pt(11)

    # Output to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


async def export_note(request: web.Request) -> web.StreamResponse:
    """GET /api/notes/{id}/export?format=pdf|docx — export a note as a file."""
    note_id = int(request.match_info["id"])
    export_format = request.query.get("format", "pdf")

    try:
        note = await db.get_note(note_id)
        if not note:
            return error_response("笔记不存在", code=404)
    except Exception as e:
        return error_response(f"获取笔记失败: {str(e)}")

    title = (note.title or "").strip() or "未命名笔记"
    content = _html_to_text(note.content or "")
    import re as _re_safe2
    safe_title = _re_safe2.sub(r'[<>:"/\\|?*()!&;$`\s]', '_', title)[:50]

    try:
        if export_format == "pdf":
            file_bytes = _export_pdf_bytes(title, content)
            filename = f"{safe_title}.pdf"
            content_type = "application/pdf"
        elif export_format == "docx":
            file_bytes = _export_docx_bytes(title, content)
            filename = f"{safe_title}.docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            return error_response(f"不支持的导出格式: {export_format}，请使用 pdf 或 docx", code=400)
    except Exception as e:
        return error_response(f"导出失败: {str(e)}")

    resp = web.StreamResponse(
        status=200,
        reason="OK",
        headers={
            "Content-Type": content_type,
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(file_bytes)),
        },
    )
    await resp.prepare(request)
    await resp.write(file_bytes)
    await resp.write_eof()
    return resp


async def share_note_to_qq(request: web.Request) -> web.Response:
    """POST /api/notes/{id}/share-to-qq — share a note as PDF/DOCX to QQ.
    
    Request body: { "format": "pdf"|"docx", "user_id": 2674610176 }
    """
    note_id = int(request.match_info["id"])

    try:
        data = await request.json()
    except json.JSONDecodeError:
        return error_response("无效的JSON数据")

    export_format = data.get("format", "pdf")
    user_id = data.get("user_id", 2674610176)
    message = (data.get("message") or "").strip()

    if export_format not in ("pdf", "docx"):
        return error_response("不支持的格式，请使用 pdf 或 docx", code=400)

    try:
        note = await db.get_note(note_id)
        if not note:
            return error_response("笔记不存在", code=404)
    except Exception as e:
        return error_response(f"获取笔记失败: {str(e)}")

    title = (note.title or "").strip() or "未命名笔记"
    content = _html_to_text(note.content or "")
    # Sanitize for filename: remove shell-dangerous chars, keep CJK/ASCII
    import re as _re_safe
    safe_title = _re_safe.sub(r'[<>:"/\\|?*()!&;$`\s]', '_', title)[:50]
    ext = "pdf" if export_format == "pdf" else "docx"

    try:
        if export_format == "pdf":
            file_bytes = _export_pdf_bytes(title, content)
        else:
            file_bytes = _export_docx_bytes(title, content)
    except Exception as e:
        return error_response(f"生成文件失败: {str(e)}")

    # Save to temp file
    import tempfile
    import os as _os
    tmp_path = _os.path.join(tempfile.gettempdir(), f"schedule_share_{note_id}_{safe_title}.{ext}")
    with open(tmp_path, "wb") as f:
        f.write(file_bytes)

    # Send to QQ
    try:
        import sys as _sys
        _sys.path.insert(0, "/home/gaoming/.opencode/skills/qq-notify")
        from send_message import send_private_message, send_private_file

        # Send text message first
        text_msg = message or f"📄 分享笔记：{title}"
        r1 = send_private_message(user_id, text_msg)

        # Send file
        file_name = f"{safe_title}.{ext}"
        r2 = send_private_file(user_id, tmp_path, file_name)

        # Clean up temp file
        _os.unlink(tmp_path)

        msg_ok = r1.get("status") == "ok"
        file_ok = r2.get("status") == "ok"
        return json_response({
            "success": msg_ok and file_ok,
            "text_sent": msg_ok,
            "file_sent": file_ok,
            "format": export_format,
        })
    except Exception as e:
        # Clean up temp file on error too
        try:
            _os.unlink(tmp_path)
        except OSError:
            pass
        return error_response(f"QQ发送失败: {str(e)}")


# ============= Route Registration =============

def register_routes(app: web.Application) -> None:
    app.router.add_get("/api/notes/{id}/export", export_note)
    app.router.add_post("/api/notes/{id}/share-to-qq", share_note_to_qq)
    app.router.add_get("/api/notes", get_notes)
    app.router.add_post("/api/notes", create_note)
    app.router.add_put("/api/notes/{id}", update_note)
    app.router.add_delete("/api/notes/{id}", delete_note)
    app.router.add_get("/api/notes/{note_id}/conversations", get_note_conversations)
    app.router.add_post("/api/notes/{note_id}/chat", chat_note)
    app.router.add_delete("/api/notes/{note_id}/conversations", delete_note_conversations)
    app.router.add_get("/api/note-groups", get_note_groups)
    app.router.add_post("/api/note-groups", create_note_group)
    app.router.add_put("/api/note-groups/{id}", update_note_group)
    app.router.add_delete("/api/note-groups/{id}", delete_note_group)
    app.router.add_put("/api/notes/reorder", reorder_notes)
